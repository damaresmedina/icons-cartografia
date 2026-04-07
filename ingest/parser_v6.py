import re
import json
import csv
import os
from collections import Counter, defaultdict
from docx import Document

BASE = os.path.join("C:", os.sep, "projetos", "icons", "ingest")
OUT = os.path.join(BASE, "saida_parser")
os.makedirs(OUT, exist_ok=True)

# ══════════════════════════════════════════
# EDITORIAIS + METADADOS
# ══════════════════════════════════════════

EDITORIAIS = [
    (re.compile(r'^Controle\s+concentrado', re.I),      'controle_concentrado'),
    (re.compile(r'^Controle\s+de\s+concentrado', re.I),  'controle_concentrado'),
    (re.compile(r'^Repercuss[aã]o\s+geral', re.I),      'repercussao_geral'),
    (re.compile(r'^Julgados?\s+[Cc]orrelatos?', re.I),  'julgados_correlatos'),
    (re.compile(r'^S[uú]mulas?\s+[Vv]inculantes?', re.I), 'sumula_vinculante'),
    (re.compile(r'^S[uú]mulas?\s*$', re.I),             'sumula'),
]

CLASSES_PROC = r'ADI|ADC|ADPF|ADO|RE|ARE|HC|RHC|MS|MI|AP|ACO|Rcl|AI|Pet|SS|SL|IF|Inq|Ext|AC|AgR|ED|MC|QO|SV|PSV|AO|AR|RMS|CR|Rp|STA'
RE_META  = re.compile(rf'^\[.*({CLASSES_PROC}).*\]$', re.IGNORECASE)
RE_META2 = re.compile(rf'^\[?\s*({CLASSES_PROC})[\s\.\-]*\d', re.IGNORECASE)
RE_SUMV  = re.compile(r'^\[?\s*S[uú]mula\s+(Vinculante\s+)?n?[oº°]?\s*\d+', re.IGNORECASE)

RE_PROC   = re.compile(rf'\b({CLASSES_PROC})\s+[\-n.]*\s*([\d.]+(?:[-/]\d+)?)', re.IGNORECASE)
RE_REL    = re.compile(r'rel\.?\s*(?:p/\s*o?\s*ac\.?\s*)?min\.?\s*([^,\]]+)', re.IGNORECASE)
RE_JULG   = re.compile(r'j\.\s*(\d[\d\-]+(?:\s*de\s*\w+\s*de\s*\d+)?)', re.IGNORECASE)
RE_DJE    = re.compile(r'D[Jj][Ee]?\s*de\s*(\d[\d\-/]+)', re.IGNORECASE)
RE_TEMA   = re.compile(r'[Tt]ema\s*(\d+)', re.IGNORECASE)

RE_INC = re.compile(r'^(I{1,3}V?|VI{0,3}|IX|X{1,3}I{0,3}|XL|L|XC|C{1,3}|D)\s*[-\u2013\u2014]\s*\S')
RE_PAR = re.compile(r'^§\s*(\d+|[uú]nico)[º°]?\s', re.IGNORECASE)
RE_ALI = re.compile(r'^([a-z])\)\s')

# ══════════════════════════════════════════
# FASE A — Árvore canônica
# ══════════════════════════════════════════

print("FASE A — Arvore canonica da CF...")
doc_cf = Document(os.path.join(BASE, "CF_1988_arvore_completa.docx"))

RE_ART_CF = re.compile(r'^Art\.?\s*(\d+)[º°]?\b')
RE_PAR_CF = re.compile(r'^§\s*(\d+|único)[º°]?\b', re.IGNORECASE)
RE_INC_CF = re.compile(r'^(I{1,3}V?|VI{0,3}|IX|X{1,3}I{0,3}|XL|L|XC|C{1,3}|D)\s*[-\u2013\u2014]')
RE_ALI_CF = re.compile(r'^([a-z])\)\s')

norm_tree = []
ctx = {'titulo': None, 'capitulo': None, 'secao': None}
cur_art = cur_par = cur_inc = None
nid = 0

for p in doc_cf.paragraphs:
    t = p.text.strip()
    if not t: continue
    s = p.style.name
    nid += 1

    if s == 'Heading 1':
        ctx = {'titulo': t, 'capitulo': None, 'secao': None}
        cur_art = cur_par = cur_inc = None
        norm_tree.append({'nid': nid, 'type': 'titulo', 'label': t, 'parent': None})
    elif s == 'Heading 2':
        ctx['capitulo'] = t; ctx['secao'] = None
        norm_tree.append({'nid': nid, 'type': 'capitulo', 'label': t, 'parent': ctx['titulo']})
    elif s == 'Heading 3':
        ctx['secao'] = t
        norm_tree.append({'nid': nid, 'type': 'secao', 'label': t, 'parent': ctx['capitulo']})
    else:
        m = RE_ART_CF.match(t)
        if m:
            cur_art = int(m.group(1)); cur_par = cur_inc = None
            norm_tree.append({'nid': nid, 'type': 'artigo', 'number': cur_art, 'label': t[:80],
                              'titulo': ctx['titulo'], 'capitulo': ctx['capitulo'], 'secao': ctx['secao']})
            continue
        m = RE_PAR_CF.match(t)
        if m and cur_art:
            cur_par = m.group(1); cur_inc = None
            norm_tree.append({'nid': nid, 'type': 'paragrafo', 'number': cur_par, 'artigo': cur_art, 'label': t[:80]})
            continue
        m = RE_INC_CF.match(t)
        if m and cur_art:
            cur_inc = m.group(1)
            norm_tree.append({'nid': nid, 'type': 'inciso', 'number': cur_inc, 'artigo': cur_art, 'label': t[:80]})
            continue
        m = RE_ALI_CF.match(t)
        if m and cur_art:
            norm_tree.append({'nid': nid, 'type': 'alinea', 'number': m.group(1), 'artigo': cur_art, 'label': t[:80]})

cf_valid_arts = set(n['number'] for n in norm_tree if n['type'] == 'artigo')

# Índices de dispositivos válidos por artigo
incisos_validos   = defaultdict(set)
paragrafos_validos = defaultdict(set)
alineas_validas   = defaultdict(set)

for no in norm_tree:
    art = no.get('artigo')
    if not art: continue
    if no.get('type') == 'inciso':
        incisos_validos[art].add(no['number'])
    elif no.get('type') == 'paragrafo':
        paragrafos_validos[art].add(str(no['number']))
    elif no.get('type') == 'alinea':
        alineas_validas[art].add(no['number'])

print(f"  Nos: {len(norm_tree)}")
print(f"  Artigos validos: {len(cf_valid_arts)}")
print(f"  Artigos com incisos: {len(incisos_validos)}")
print(f"  Artigos com paragrafos: {len(paragrafos_validos)}")

# ══════════════════════════════════════════
# FASE B — Classificador
# ══════════════════════════════════════════

print("\nFASE B — Classificando paragrafos...")
doc_main = Document(os.path.join(BASE, "cf_comentada.docx"))

falsos_positivos = 0
artigos_vistos = set()
ultimo_artigo_cf = 0

def e_ancora_cf(texto):
    global falsos_positivos, ultimo_artigo_cf
    t = texto.strip()
    m = re.match(r'^Art\.?\s*(\d+)[º°]?', t)
    if not m: return False, None
    num = int(m.group(1))
    if num < 1 or num > 250: return False, None
    if num not in cf_valid_arts: return False, None

    resto = t[m.end():m.end()+120]
    indicadores = [
        r'^\s*,\s*§',                          # Art. 10, § 5º
        r'^\s*,\s*[Cc]aput',                   # Art. 55, caput
        r'^\s*,\s*[Ii]ncisos?',                # Art. 55, incisos
        r'^\s*,\s*[Pp]ar[aá]grafo\s+[uú]nico', # Art. 137, parágrafo único, da Lei
        r'^\s*,\s*\d+[º°]',                    # Art. 10, 2º
        r'^\s+e\s+art\.',                       # Art. 5 e art. 10
        r'^\s*\/\s*art\.',                      # Art. 5/art. 10
        r'^\s+da Lei\b',
        r'^\s+do C[oó]digo',
        r'^\s+da CLT',
        r'^\s+do CPC\b',
        r'^\s+do CP\b',
        r'^\s+do CTN',
        r'^\s+da LC\b',
        r'^\s+Lei Complementar',
        r'^\s+da Constitui[çc][ãa]o do Estado',
        r'^\s+da Lei Org[âa]nica',
        r'^\s+do Estatuto',
        r'^\s+do Decreto',
        r'^\s+da Resolu[çc][ãa]o',
        r'^\s+c/c\b',
        r'^\s+do CPP\b',
        r'^\s+do RISTF\b',
        r'^\s+da LEP\b',
        r'^\s+da LOMAN\b',
        r'^\s+Decreto-lei',
    ]
    for ind in indicadores:
        if re.match(ind, resto, re.IGNORECASE):
            falsos_positivos += 1
            return False, None

    if num in artigos_vistos:
        falsos_positivos += 1
        return False, None

    artigos_vistos.add(num)
    ultimo_artigo_cf = num
    return True, num

def classifica(texto):
    t = texto.strip()

    is_ancora, art_num = e_ancora_cf(t)
    if is_ancora:
        return ('A', 'artigo', art_num)

    m = RE_PAR.match(t)
    if m and len(t) < 300:
        return ('A', 'paragrafo', m.group(1).lower().replace('\u00fa', 'u'))

    m = re.match(r'^Par[a\u00e1]grafo\s+[u\u00fa]nico', t, re.IGNORECASE)
    if m and len(t) < 300:
        return ('A', 'paragrafo', 'unico')

    m = RE_INC.match(t)
    if m and len(t) < 200:
        return ('A', 'inciso', m.group(1))

    m = RE_ALI.match(t)
    if m and len(t) < 200:
        return ('A', 'alinea', m.group(1))

    for pat, cat in EDITORIAIS:
        if pat.match(t): return ('B', cat, None)

    if RE_META.match(t) or RE_META2.match(t) or RE_SUMV.match(t):
        return ('D', None, None)

    return ('C', None, None)

paragrafos = []
for i, p in enumerate(doc_main.paragraphs):
    t = p.text.strip()
    if not t: continue
    cl, subtipo, valor = classifica(t)
    paragrafos.append({'idx': i, 'classe': cl, 'subtipo': subtipo, 'valor': valor, 'texto': t})

dist = Counter(p['classe'] for p in paragrafos)
print(f"  Total: {len(paragrafos)}")
print(f"  Distribuicao: {dict(dist)}")

# ══════════════════════════════════════════
# FASE C — Máquina de estados com validação
# ══════════════════════════════════════════

print("\nFASE C — Maquina de estados com validacao...")

def build_slug(state):
    art = state['artigo']
    if not art: return None
    slug = f'cf-art-{art}'
    if state['paragrafo']: slug += f'-par-{state["paragrafo"]}'
    if state['inciso']:    slug += f'-inc-{state["inciso"]}'
    if state['alinea']:    slug += f'-ali-{state["alinea"]}'
    return slug

def extract_meta(text):
    meta = {'process_class': '', 'process_number': '', 'relator': '', 'data_julgamento': '', 'dje': '', 'tema': ''}
    m = RE_PROC.search(text)
    if m: meta['process_class'] = m.group(1).upper(); meta['process_number'] = m.group(2)
    m = RE_REL.search(text)
    if m: meta['relator'] = m.group(1).strip()
    m = RE_JULG.search(text)
    if m: meta['data_julgamento'] = m.group(1).strip()
    m = RE_DJE.search(text)
    if m: meta['dje'] = m.group(1).strip()
    m = RE_TEMA.search(text)
    if m: meta['tema'] = m.group(1)
    return meta

state = {'artigo': None, 'paragrafo': None, 'inciso': None, 'alinea': None, 'editorial': 'sem_categoria', 'slug': None}
blocks = []
block_id = 0
current_block = None
rejected_subnodes = 0

def flush():
    global current_block
    if current_block and (current_block['block_text'].strip() or current_block['metadata_text'].strip()):
        merged = (current_block['block_text'] + ' ' + current_block['metadata_text']).strip()
        meta = extract_meta(merged)
        current_block.update(meta)
        blocks.append(current_block)
    current_block = None

def new_block(idx):
    global block_id
    block_id += 1
    return {
        'block_id': block_id,
        'anchor_slug': state['slug'],
        'anchor_artigo': state['artigo'],
        'anchor_paragrafo': state['paragrafo'],
        'anchor_inciso': state['inciso'],
        'anchor_alinea': state['alinea'],
        'editorial_marker': state['editorial'],
        'block_text': '',
        'metadata_text': '',
        'paragraph_start': idx,
        'paragraph_end': idx
    }

def add_text(idx, texto):
    global current_block
    if not current_block:
        current_block = new_block(idx)
    current_block['block_text'] += (' ' if current_block['block_text'] else '') + texto
    current_block['paragraph_end'] = idx

for p in paragrafos:
    idx = p['idx']
    cl = p['classe']
    subtipo = p['subtipo']
    valor = p['valor']
    texto = p['texto']

    if cl == 'A':
        if subtipo == 'artigo':
            flush()
            state['artigo']    = valor
            state['paragrafo'] = None
            state['inciso']    = None
            state['alinea']    = None
            state['editorial'] = 'sem_categoria'
            state['slug'] = build_slug(state)

        elif subtipo == 'paragrafo' and state['artigo']:
            par_num = str(valor)
            if par_num in paragrafos_validos.get(state['artigo'], set()) or not paragrafos_validos.get(state['artigo']):
                flush()
                state['paragrafo'] = valor
                state['inciso']    = None
                state['alinea']    = None
                state['slug'] = build_slug(state)
            else:
                rejected_subnodes += 1
                add_text(idx, texto)

        elif subtipo == 'inciso' and state['artigo']:
            inc_num = valor
            if inc_num in incisos_validos.get(state['artigo'], set()) or not incisos_validos.get(state['artigo']):
                flush()
                state['inciso']  = valor
                state['alinea']  = None
                state['slug'] = build_slug(state)
            else:
                rejected_subnodes += 1
                add_text(idx, texto)

        elif subtipo == 'alinea' and state['artigo']:
            ali_num = valor
            if ali_num in alineas_validas.get(state['artigo'], set()) or not alineas_validas.get(state['artigo']):
                flush()
                state['alinea'] = valor
                state['slug'] = build_slug(state)
            else:
                rejected_subnodes += 1
                add_text(idx, texto)
        else:
            add_text(idx, texto)

    elif cl == 'B':
        flush()
        state['editorial'] = subtipo

    elif cl == 'C':
        add_text(idx, texto)

    elif cl == 'D':
        if not current_block:
            current_block = new_block(idx)
        current_block['metadata_text'] += (' ' if current_block['metadata_text'] else '') + texto
        current_block['paragraph_end'] = idx
        flush()

flush()
print(f"  Blocos: {len(blocks)}")
print(f"  Sub-nos rejeitados (nao existem na arvore): {rejected_subnodes}")

# ══════════════════════════════════════════
# DIAGNÓSTICO
# ══════════════════════════════════════════

arts_cobertos = set(b['anchor_artigo'] for b in blocks if b['anchor_artigo'])
faltando = sorted(set(range(1, 251)) - arts_cobertos)

print(f"\n=== DIAGNOSTICO FINAL ===")
print(f"  Nos arvore: {len(norm_tree)}")
print(f"  Blocos: {len(blocks)}")
print(f"  Artigos cobertos: {len(arts_cobertos)}")
print(f"  Faltando ({len(faltando)}): {faltando}")
print(f"  Falsos positivos Art.: {falsos_positivos}")
print(f"  Sub-nos rejeitados: {rejected_subnodes}")

print(f"\n  Blocos com inciso: {sum(1 for b in blocks if b['anchor_inciso'])}")
print(f"  Blocos com paragrafo: {sum(1 for b in blocks if b['anchor_paragrafo'])}")
print(f"  Blocos com alinea: {sum(1 for b in blocks if b['anchor_alinea'])}")

# Top 20 slugs
slug_counts = Counter(b['anchor_slug'] for b in blocks if b['anchor_slug'])
print(f"\n  TOP 20 slugs por blocos:")
for slug, cnt in slug_counts.most_common(20):
    print(f"    {slug}: {cnt}")

# Ramificação Art. 5
art5 = [b for b in blocks if b['anchor_artigo'] == 5]
inc5 = Counter(b.get('anchor_inciso') or 'raiz' for b in art5)
print(f"\n  Art. 5: {len(art5)} blocos")
print(f"    Ramos (top 10):")
for inc, cnt in inc5.most_common(10):
    print(f"      inc-{inc}: {cnt}")

# Ramificação Art. 37
art37 = [b for b in blocks if b['anchor_artigo'] == 37]
inc37 = Counter(b.get('anchor_inciso') or 'raiz' for b in art37)
print(f"\n  Art. 37: {len(art37)} blocos")
print(f"    Ramos (top 10):")
for inc, cnt in inc37.most_common(10):
    print(f"      inc-{inc}: {cnt}")

# Ramificação Art. 102
art102 = [b for b in blocks if b['anchor_artigo'] == 102]
inc102 = Counter(b.get('anchor_inciso') or 'raiz' for b in art102)
print(f"\n  Art. 102: {len(art102)} blocos")
print(f"    Ramos (top 10):")
for inc, cnt in inc102.most_common(10):
    print(f"      inc-{inc}: {cnt}")

# Editorial
eds = Counter(b['editorial_marker'] for b in blocks)
print(f"\n  Por editorial:")
for ed, n in eds.most_common():
    print(f"    {ed}: {n}")

procs = sum(1 for b in blocks if b['process_class'])
print(f"\n  Blocos com processo: {procs}")

# ══════════════════════════════════════════
# SALVAR
# ══════════════════════════════════════════

with open(os.path.join(OUT, 'norm_tree.json'), 'w', encoding='utf-8') as f:
    json.dump(norm_tree, f, ensure_ascii=False, indent=2)

with open(os.path.join(OUT, 'commentary_blocks.json'), 'w', encoding='utf-8') as f:
    json.dump(blocks, f, ensure_ascii=False, indent=2)

with open(os.path.join(OUT, 'commentary_blocks.tsv'), 'w', encoding='utf-8-sig', newline='') as f:
    if blocks:
        w = csv.DictWriter(f, fieldnames=list(blocks[0].keys()), delimiter='\t')
        w.writeheader()
        w.writerows(blocks)

print(f"\n  Arquivos em: {OUT}")
