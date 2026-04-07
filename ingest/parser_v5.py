import re
import json
import csv
import os
from collections import Counter
from docx import Document

BASE = os.path.join("C:", os.sep, "projetos", "icons", "ingest")
OUT = os.path.join(BASE, "saida_parser")
os.makedirs(OUT, exist_ok=True)

# ══════════════════════════════════════════
# EDITORIAIS + METADADOS
# ══════════════════════════════════════════

EDITORIAIS = [
    (re.compile(r'^Controle\s+concentrado', re.I),      'controle_concentrado'),
    (re.compile(r'^Controle\s+de\s+concentrado', re.I),  'controle_concentrado'),
    (re.compile(r'^Repercuss[aã]o\s+geral', re.I),      'repercussao_geral'),
    (re.compile(r'^Julgados?\s+[Cc]orrelatos?', re.I),  'julgados_correlatos'),
    (re.compile(r'^S[uú]mulas?\s+[Vv]inculantes?', re.I), 'sumula_vinculante'),
    (re.compile(r'^S[uú]mulas?\s*$', re.I),             'sumula'),
]

CLASSES_PROC = r'ADI|ADC|ADPF|ADO|RE|ARE|HC|RHC|MS|MI|AP|ACO|Rcl|AI|Pet|SS|SL|IF|Inq|Ext|AC|AgR|ED|MC|QO|SV|PSV|AO|AR|RMS|CR|Rp|STA'
RE_META  = re.compile(rf'^\[.*({CLASSES_PROC}).*\]$', re.IGNORECASE)
RE_META2 = re.compile(rf'^\[?\s*({CLASSES_PROC})[\s\.\-]*\d', re.IGNORECASE)
RE_SUMV  = re.compile(r'^\[?\s*S[uú]mula\s+(Vinculante\s+)?n?[oº°]?\s*\d+', re.IGNORECASE)

# Process metadata extraction
RE_PROC   = re.compile(rf'\b({CLASSES_PROC})\s+[\-n.]*\s*([\d.]+(?:[-/]\d+)?)', re.IGNORECASE)
RE_REL    = re.compile(r'rel\.?\s*(?:p/\s*o?\s*ac\.?\s*)?min\.?\s*([^,\]]+)', re.IGNORECASE)
RE_JULG   = re.compile(r'j\.\s*(\d[\d\-]+(?:\s*de\s*\w+\s*de\s*\d+)?)', re.IGNORECASE)
RE_DJE    = re.compile(r'D[Jj][Ee]?\s*de\s*(\d[\d\-/]+)', re.IGNORECASE)
RE_TEMA   = re.compile(r'[Tt]ema\s*(\d+)', re.IGNORECASE)

# ══════════════════════════════════════════
# FASE A — Árvore canônica
# ══════════════════════════════════════════

print("FASE A — Arvore canonica da CF...")
doc_cf = Document(os.path.join(BASE, "CF_1988_arvore_completa.docx"))

RE_ART_CF = re.compile(r'^Art\.?\s*(\d+)[º°]?\b')
RE_PAR_CF = re.compile(r'^§\s*(\d+|único)[º°]?\b', re.IGNORECASE)
RE_INC_CF = re.compile(r'^(I{1,3}V?|VI{0,3}|IX|X{1,3}I{0,3}|XL|L|XC|C)\s*[-\u2013\u2014]')
RE_ALI_CF = re.compile(r'^([a-z])\)\s')

norm_tree = []
ctx = {'titulo': None, 'capitulo': None, 'secao': None}
cur_art = cur_par = cur_inc = None
nid = 0

for p in doc_cf.paragraphs:
    t = p.text.strip()
    if not t: continue
    s = p.style.name
    nid += 1

    if s == 'Heading 1':
        ctx = {'titulo': t, 'capitulo': None, 'secao': None}
        cur_art = cur_par = cur_inc = None
        norm_tree.append({'nid': nid, 'type': 'titulo', 'label': t, 'parent': None})
    elif s == 'Heading 2':
        ctx['capitulo'] = t; ctx['secao'] = None
        norm_tree.append({'nid': nid, 'type': 'capitulo', 'label': t, 'parent': ctx['titulo']})
    elif s == 'Heading 3':
        ctx['secao'] = t
        norm_tree.append({'nid': nid, 'type': 'secao', 'label': t, 'parent': ctx['capitulo']})
    else:
        m = RE_ART_CF.match(t)
        if m:
            cur_art = int(m.group(1)); cur_par = cur_inc = None
            norm_tree.append({'nid': nid, 'type': 'artigo', 'number': cur_art, 'label': t[:80],
                              'titulo': ctx['titulo'], 'capitulo': ctx['capitulo'], 'secao': ctx['secao']})
            continue
        m = RE_PAR_CF.match(t)
        if m and cur_art:
            cur_par = m.group(1); cur_inc = None
            norm_tree.append({'nid': nid, 'type': 'paragrafo', 'number': cur_par, 'artigo': cur_art, 'label': t[:80]})
            continue
        m = RE_INC_CF.match(t)
        if m and cur_art:
            cur_inc = m.group(1)
            norm_tree.append({'nid': nid, 'type': 'inciso', 'number': cur_inc, 'artigo': cur_art, 'label': t[:80]})
            continue
        m = RE_ALI_CF.match(t)
        if m and cur_art:
            norm_tree.append({'nid': nid, 'type': 'alinea', 'number': m.group(1), 'artigo': cur_art, 'label': t[:80]})

# Valid CF article numbers
cf_valid_arts = set(n['number'] for n in norm_tree if n['type'] == 'artigo')

# Build valid sub-nodes per article from norm_tree
from collections import defaultdict
incisos_validos = defaultdict(set)
paragrafos_validos = defaultdict(set)
alineas_validas = defaultdict(set)

for no in norm_tree:
    if no.get('type') == 'inciso' and no.get('artigo'):
        incisos_validos[no['artigo']].add(no['number'])
    elif no.get('type') == 'paragrafo' and no.get('artigo'):
        paragrafos_validos[no['artigo']].add(str(no['number']))
    elif no.get('type') == 'alinea' and no.get('artigo'):
        alineas_validas[no['artigo']].add(no['number'])

print(f"  Nos: {len(norm_tree)}")
print(f"  Artigos validos: {len(cf_valid_arts)}")
print(f"  Artigos com incisos: {len(incisos_validos)}")
print(f"  Artigos com paragrafos: {len(paragrafos_validos)}")

# ══════════════════════════════════════════
# FASE B — Classificador A/B/C/D
# ══════════════════════════════════════════

print("\nFASE B — Classificando paragrafos...")
doc_main = Document(os.path.join(BASE, "cf_comentada.docx"))

# Indicadores de referência cruzada (Art. de outra lei, não da CF)
INDICADORES_REF = [
    re.compile(r'da Lei', re.I),
    re.compile(r'do C[oó]digo', re.I),
    re.compile(r'da CLT', re.I),
    re.compile(r'do CPC', re.I),
    re.compile(r'do CP\b', re.I),
    re.compile(r'do CTN', re.I),
    re.compile(r'da LC\b', re.I),
    re.compile(r'Lei Complementar', re.I),
    re.compile(r'da Constitui[çc][ãa]o do Estado', re.I),
    re.compile(r'da Lei Org[âa]nica', re.I),
    re.compile(r'do Estatuto', re.I),
    re.compile(r'do Decreto', re.I),
    re.compile(r'da Resolu[çc][ãa]o', re.I),
    re.compile(r'da Portaria', re.I),
    re.compile(r'c/c\b', re.I),
    re.compile(r'combinado com', re.I),
    re.compile(r'nos termos do', re.I),
    re.compile(r'conforme o', re.I),
    re.compile(r'previsto no', re.I),
    re.compile(r'disposto no', re.I),
    re.compile(r'do RISTF', re.I),
    re.compile(r'da LEP', re.I),
    re.compile(r'da LOMAN', re.I),
    re.compile(r'do CPP', re.I),
    re.compile(r'Lei n', re.I),
    re.compile(r'Decreto-lei', re.I),
]

falsos_positivos_eliminados = 0
artigos_vistos = set()

def e_ancora_cf(texto):
    """Distingue âncora CF de referência cruzada a outra lei.
    Só aceita a primeira ocorrência de cada artigo no documento."""
    global falsos_positivos_eliminados
    t = texto.strip()

    m = re.match(r'^Art\.?\s*(\d+)[º°]?', t)
    if not m:
        return False, None

    num = int(m.group(1))

    if num < 1 or num > 250:
        return False, None

    if num not in cf_valid_arts:
        return False, None

    pos = m.end()
    resto = t[pos:pos+120]

    # Referência cruzada — indicadores IMEDIATOS após o número
    indicadores = [
        r'^\s*,\s*§',           # Art. 10, § 5º
        r'^\s*,\s*[Cc]aput',    # Art. 55, caput
        r'^\s*,\s*[Ii]ncisos?', # Art. 55, incisos
        r'^\s+da Lei\b',
        r'^\s+do C[oó]digo',
        r'^\s+da CLT',
        r'^\s+do CPC\b',
        r'^\s+do CP\b',
        r'^\s+do CTN',
        r'^\s+da LC\b',
        r'^\s+Lei Complementar',
        r'^\s+da Constitui[çc][ãa]o do Estado',
        r'^\s+da Lei Org[âa]nica',
        r'^\s+do Estatuto',
        r'^\s+do Decreto',
        r'^\s+da Resolu[çc][ãa]o',
        r'^\s+c/c\b',
        r'^\s+do CPP\b',
        r'^\s+do RISTF\b',
        r'^\s+da LEP\b',
        r'^\s+da LOMAN\b',
        r'^\s+Decreto-lei',
    ]
    for ind in indicadores:
        if re.match(ind, resto, re.IGNORECASE):
            falsos_positivos_eliminados += 1
            return False, None

    # Só aceita primeira ocorrência de cada artigo
    if num in artigos_vistos:
        falsos_positivos_eliminados += 1
        return False, None

    artigos_vistos.add(num)
    return True, num

def classifica(texto):
    t = texto.strip()

    # 1. Verifica se é âncora CF válida (não referência cruzada)
    is_ancora, art_num = e_ancora_cf(t)
    if is_ancora:
        return ('A', 'artigo', art_num)

    # 2. Parágrafo normativo
    m = re.match(r'^§\s*(\d+|[uú]nico)[º°]?\s', t, re.IGNORECASE)
    if m and len(t) < 300:
        return ('A', 'paragrafo', m.group(1).lower().replace('\u00fa', 'u'))

    # Parágrafo único alternativo
    m = re.match(r'^Par[a\u00e1]grafo\s+[u\u00fa]nico', t, re.IGNORECASE)
    if m and len(t) < 300:
        return ('A', 'paragrafo', 'unico')

    # 3. Inciso normativo — curto e começa com romano
    m = re.match(r'^(I{1,3}V?|VI{0,3}|IX|X{1,3}I{0,3}|XL|L|XC|C{1,3}|D)\s*[-\u2013\u2014]\s*\S', t)
    if m and len(t) < 200:
        return ('A', 'inciso', m.group(1))

    # 4. Alínea normativa — curta
    m = re.match(r'^([a-z])\)\s', t)
    if m and len(t) < 200:
        return ('A', 'alinea', m.group(1))

    # 5. Editorial
    for pat, cat in EDITORIAIS:
        if pat.match(t): return ('B', cat, None)

    # 6. Metadado
    if RE_META.match(t) or RE_META2.match(t) or RE_SUMV.match(t):
        return ('D', None, None)

    return ('C', None, None)

paragrafos = []
for i, p in enumerate(doc_main.paragraphs):
    t = p.text.strip()
    if not t: continue
    cl, subtipo, valor = classifica(t)
    paragrafos.append({'idx': i, 'classe': cl, 'subtipo': subtipo, 'valor': valor, 'texto': t})

dist = Counter(p['classe'] for p in paragrafos)
print(f"  Total: {len(paragrafos)}")
print(f"  Distribuicao: {dict(dist)}")

# ══════════════════════════════════════════
# FASE C — Máquina de estados
# ══════════════════════════════════════════

print("\nFASE C — Maquina de estados...")

def build_slug(state):
    art = state['artigo']
    if not art: return None
    slug = f'cf-art-{art}'
    if state['paragrafo']: slug += f'-par-{state["paragrafo"]}'
    if state['inciso']:    slug += f'-inc-{state["inciso"]}'
    if state['alinea']:    slug += f'-ali-{state["alinea"]}'
    return slug

def atualiza_state(subtipo, valor, state):
    """Atualiza state somente se o nó existe na norm_tree para o artigo ativo."""
    if subtipo == 'artigo':
        state['artigo']    = valor
        state['paragrafo'] = None
        state['inciso']    = None
        state['alinea']    = None
        state['editorial'] = 'sem_categoria'
    elif subtipo == 'paragrafo':
        # Só aceita se parágrafo existe para este artigo na árvore
        art = state['artigo']
        if art and str(valor) in paragrafos_validos.get(art, set()):
            state['paragrafo'] = valor
            state['inciso']    = None
            state['alinea']    = None
        else:
            # Parágrafo não existe na árvore — não atualiza, trata como texto
            state['slug'] = build_slug(state)
            return False
    elif subtipo == 'inciso':
        # Só aceita se inciso existe para este artigo na árvore
        art = state['artigo']
        if art and valor in incisos_validos.get(art, set()):
            state['inciso']  = valor
            state['alinea']  = None
        else:
            # Inciso não existe na árvore — não atualiza, trata como texto
            state['slug'] = build_slug(state)
            return False
    elif subtipo == 'alinea':
        # Só aceita se alínea existe para este artigo na árvore
        art = state['artigo']
        if art and valor in alineas_validas.get(art, set()):
            state['alinea'] = valor
        else:
            state['slug'] = build_slug(state)
            return False
    state['slug'] = build_slug(state)
    return True

def extract_meta(text):
    meta = {'process_class': '', 'process_number': '', 'relator': '', 'data_julgamento': '', 'dje': '', 'tema': ''}
    m = RE_PROC.search(text)
    if m: meta['process_class'] = m.group(1).upper(); meta['process_number'] = m.group(2)
    m = RE_REL.search(text)
    if m: meta['relator'] = m.group(1).strip()
    m = RE_JULG.search(text)
    if m: meta['data_julgamento'] = m.group(1).strip()
    m = RE_DJE.search(text)
    if m: meta['dje'] = m.group(1).strip()
    m = RE_TEMA.search(text)
    if m: meta['tema'] = m.group(1)
    return meta

state = {'artigo': None, 'paragrafo': None, 'inciso': None, 'alinea': None, 'editorial': 'sem_categoria', 'slug': None}
blocks = []
block_id = 0
current_block = None

def flush():
    global current_block
    if current_block and (current_block['block_text'].strip() or current_block['metadata_text'].strip()):
        merged = (current_block['block_text'] + ' ' + current_block['metadata_text']).strip()
        meta = extract_meta(merged)
        current_block.update(meta)
        blocks.append(current_block)
    current_block = None

def new_block(idx):
    global block_id
    block_id += 1
    return {
        'block_id': block_id,
        'anchor_slug': state['slug'],
        'anchor_artigo': state['artigo'],
        'anchor_paragrafo': state['paragrafo'],
        'anchor_inciso': state['inciso'],
        'anchor_alinea': state['alinea'],
        'editorial_marker': state['editorial'],
        'block_text': '',
        'metadata_text': '',
        'paragraph_start': idx,
        'paragraph_end': idx
    }

for p in paragrafos:
    idx = p['idx']
    cl = p['classe']
    subtipo = p['subtipo']
    valor = p['valor']
    texto = p['texto']

    if cl == 'A':
        if subtipo == 'artigo':
            flush()
            atualiza_state(subtipo, valor, state)
            continue
        else:
            # Inciso/parágrafo/alínea — só atualiza se válido na árvore
            accepted = atualiza_state(subtipo, valor, state)
            if accepted:
                flush()
                continue
            # Não existe na árvore — cai para tratamento como texto abaixo

    if cl == 'B':
        flush()
        state['editorial'] = subtipo

    elif cl == 'C' or (cl == 'A'):
        # cl == 'A' que caiu aqui = sub-nó rejeitado, tratar como texto
        if not current_block:
            current_block = new_block(idx)
        current_block['block_text'] += (' ' if current_block['block_text'] else '') + texto
        current_block['paragraph_end'] = idx

    elif cl == 'D':
        if not current_block:
            current_block = new_block(idx)
        current_block['metadata_text'] += (' ' if current_block['metadata_text'] else '') + texto
        current_block['paragraph_end'] = idx
        flush()

flush()
print(f"  Blocos: {len(blocks)}")

# ══════════════════════════════════════════
# DIAGNÓSTICO
# ══════════════════════════════════════════

arts_cobertos = set(b['anchor_artigo'] for b in blocks if b['anchor_artigo'])
faltando = sorted(set(range(1, 251)) - arts_cobertos)

print(f"\n=== DIAGNOSTICO FINAL ===")
print(f"  Nos arvore: {len(norm_tree)}")
print(f"  Blocos: {len(blocks)}")
print(f"  Artigos cobertos: {len(arts_cobertos)}")
print(f"  Faltando ({len(faltando)}): {faltando}")

# Blocos com slug de inciso
inc_blocks = sum(1 for b in blocks if b['anchor_inciso'])
par_blocks = sum(1 for b in blocks if b['anchor_paragrafo'])
ali_blocks = sum(1 for b in blocks if b['anchor_alinea'])
print(f"\n  Blocos com inciso: {inc_blocks}")
print(f"  Blocos com paragrafo: {par_blocks}")
print(f"  Blocos com alinea: {ali_blocks}")

# Exemplos
print(f"\n  Exemplos de slugs com inciso:")
for b in blocks:
    if b['anchor_inciso'] and b['anchor_artigo'] == 5:
        print(f"    {b['anchor_slug']} | {b['block_text'][:60]}")
        break
for b in blocks:
    if b['anchor_inciso'] and b['anchor_artigo'] == 37:
        print(f"    {b['anchor_slug']} | {b['block_text'][:60]}")
        break

print(f"\n  Exemplos de slugs com paragrafo:")
for b in blocks:
    if b['anchor_paragrafo'] and b['anchor_artigo'] == 37:
        print(f"    {b['anchor_slug']} | {b['block_text'][:60]}")
        break

print(f"\n  Exemplos de slugs com alinea:")
for b in blocks:
    if b['anchor_alinea'] and b['anchor_artigo'] == 7:
        print(f"    {b['anchor_slug']} | {b['block_text'][:60]}")
        break

# Editorial
eds = Counter(b['editorial_marker'] for b in blocks)
print(f"\n  Por editorial:")
for ed, n in eds.most_common():
    print(f"    {ed}: {n}")

# Top 10 artigos
art_counts = Counter(b['anchor_artigo'] for b in blocks if b['anchor_artigo'])
print(f"\n  TOP 10 artigos por blocos:")
for art, cnt in art_counts.most_common(10):
    print(f"    Art. {art:3d}: {cnt:4d} blocos")

# Art. 5
art5 = [b for b in blocks if b['anchor_artigo'] == 5]
print(f"\n  Art. 5: {len(art5)} blocos")
for b in art5[:3]:
    print(f"    [{b['block_id']}] [{b['anchor_slug']}] [{b['editorial_marker']}]")
    print(f"      texto: {b['block_text'][:80]}")
    print(f"      meta:  {b['metadata_text'][:80]}")

# Blocos com processo
procs = sum(1 for b in blocks if b['process_class'])
print(f"\n  Blocos com processo: {procs}")
print(f"  Falsos positivos eliminados: {falsos_positivos_eliminados}")

# ══════════════════════════════════════════
# SALVAR
# ══════════════════════════════════════════

with open(os.path.join(OUT, 'norm_tree.json'), 'w', encoding='utf-8') as f:
    json.dump(norm_tree, f, ensure_ascii=False, indent=2)

with open(os.path.join(OUT, 'commentary_blocks.json'), 'w', encoding='utf-8') as f:
    json.dump(blocks, f, ensure_ascii=False, indent=2)

with open(os.path.join(OUT, 'commentary_blocks.tsv'), 'w', encoding='utf-8-sig', newline='') as f:
    if blocks:
        w = csv.DictWriter(f, fieldnames=list(blocks[0].keys()), delimiter='\t')
        w.writeheader()
        w.writerows(blocks)

print(f"\n  Arquivos em: {OUT}")
