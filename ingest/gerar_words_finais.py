import json, os, re
from collections import Counter, defaultdict
from docx import Document
from docx.shared import Pt, RGBColor, Inches

OUT = os.path.join("C:", os.sep, "projetos", "icons", "ingest", "saida_parser")
BASE = os.path.join("C:", os.sep, "projetos", "icons", "ingest")

with open(os.path.join(OUT, "norm_tree.json"), "r", encoding="utf-8") as f:
    norm_tree = json.load(f)
with open(os.path.join(OUT, "commentary_blocks.json"), "r", encoding="utf-8") as f:
    blocks = json.load(f)

arts_cobertos = set(b['anchor_artigo'] for b in blocks if b['anchor_artigo'])
por_artigo = Counter(b['anchor_artigo'] for b in blocks if b['anchor_artigo'])
art_ramos = defaultdict(set)
for b in blocks:
    if b.get('anchor_artigo') and b.get('anchor_slug'):
        art_ramos[b['anchor_artigo']].add(b['anchor_slug'])

# ══════════════════════════════════════════
# 1. CF COM ANCORAGEM
# ══════════════════════════════════════════
print("1. ESPELHO_CF_ancoragem_v8_final.docx...")
doc1 = Document()
doc1.styles['Normal'].font.name = 'Arial'
doc1.styles['Normal'].font.size = Pt(9)

doc1.add_heading('CF 1988 — ANCORAGEM DE DECISOES', level=0)
doc1.add_paragraph(f'{len(arts_cobertos)} artigos cobertos | 9.042 blocos | Parser v8')

for node in norm_tree:
    nt = node['type']
    label = node.get('label', '')
    art_num = node.get('number', node.get('artigo', ''))

    if nt == 'titulo':
        doc1.add_heading(label, level=1)
    elif nt == 'capitulo':
        doc1.add_heading(label, level=2)
    elif nt == 'secao':
        doc1.add_heading(label, level=3)
    elif nt == 'artigo':
        dec = por_artigo.get(art_num, 0)
        ramos = len(art_ramos.get(art_num, set()))
        p = doc1.add_paragraph()
        r = p.add_run(label)
        r.bold = True
        r.font.size = Pt(10)
        if dec > 0:
            r2 = p.add_run(f'  [{dec} decisoes, {ramos} ramos]')
            r2.font.size = Pt(7)
            r2.font.color.rgb = RGBColor(0, 100, 180)
    elif nt == 'paragrafo':
        p = doc1.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        p.add_run(label).font.size = Pt(9)
    elif nt == 'inciso':
        p = doc1.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        p.add_run(label).font.size = Pt(9)
    elif nt == 'alinea':
        p = doc1.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.7)
        p.add_run(label).font.size = Pt(9)

doc1.save(os.path.join(BASE, "ESPELHO_CF_ancoragem_v8_final.docx"))
print("  Salvo.")

# ══════════════════════════════════════════
# 2. BLOCOS DE DECISOES
# ══════════════════════════════════════════
print("2. ESPELHO_blocos_decisoes_v8_final.docx...")
doc2 = Document()
doc2.styles['Normal'].font.name = 'Arial'
doc2.styles['Normal'].font.size = Pt(9)

doc2.add_heading('BLOCOS DE DECISOES E COMENTARIOS', level=0)
doc2.add_paragraph(f'{len(blocks)} blocos | Parser v8')

current_art = None
for b in blocks:
    art = b.get('anchor_artigo')
    slug = b.get('anchor_slug', '')
    editorial = b.get('editorial_marker', '')
    texto = b.get('block_text', '').strip()
    meta = b.get('metadata_text', '').strip()
    proc = b.get('process_class', '')
    num = b.get('process_number', '')
    relator = b.get('relator', '')
    ano = b.get('ano', '')

    if art != current_art:
        current_art = art
        if art:
            doc2.add_heading(f'Art. {art}', level=1)
        elif slug == 'adct':
            doc2.add_heading('ADCT', level=1)
        else:
            doc2.add_heading('Sem ancora', level=1)

    if editorial and editorial != 'sem_categoria':
        p_ed = doc2.add_paragraph()
        r_ed = p_ed.add_run(editorial.upper().replace('_', ' '))
        r_ed.bold = True
        r_ed.font.size = Pt(8)
        r_ed.font.color.rgb = RGBColor(180, 100, 0)

    p_h = doc2.add_paragraph()
    r_h = p_h.add_run(f'{slug}')
    r_h.font.size = Pt(7)
    r_h.font.color.rgb = RGBColor(140, 140, 140)
    if proc:
        r_p = p_h.add_run(f' | {proc} {num}')
        r_p.font.size = Pt(7)
        r_p.font.color.rgb = RGBColor(0, 80, 160)
    if relator:
        r_r = p_h.add_run(f' | rel. {relator}')
        r_r.font.size = Pt(7)
        r_r.font.color.rgb = RGBColor(0, 80, 160)
    if ano:
        r_a = p_h.add_run(f' | {ano}')
        r_a.font.size = Pt(7)
        r_a.font.color.rgb = RGBColor(0, 80, 160)

    if texto:
        p_t = doc2.add_paragraph()
        p_t.paragraph_format.left_indent = Inches(0.15)
        p_t.add_run(texto[:500]).font.size = Pt(9)

    if meta:
        p_m = doc2.add_paragraph()
        p_m.paragraph_format.left_indent = Inches(0.15)
        r_m = p_m.add_run(meta[:200])
        r_m.font.size = Pt(8)
        r_m.font.color.rgb = RGBColor(0, 60, 140)
        r_m.bold = True

doc2.save(os.path.join(BASE, "ESPELHO_blocos_decisoes_v8_final.docx"))
print("  Salvo.")

# ══════════════════════════════════════════
# 3. DIAGNOSTICO
# ══════════════════════════════════════════
print("3. DIAGNOSTICO_PARSER_v8_final.docx...")
doc3 = Document()
doc3.styles['Normal'].font.name = 'Arial'
doc3.styles['Normal'].font.size = Pt(10)

doc3.add_heading('DIAGNOSTICO FINAL — PARSER v8', level=0)

faltando = sorted(set(range(1, 251)) - arts_cobertos)
por_editorial = Counter(b['editorial_marker'] for b in blocks)
rels = Counter(b.get('relator') for b in blocks if b.get('relator'))
cls_proc = Counter(b.get('process_class') for b in blocks if b.get('process_class'))
adct_blocks = [b for b in blocks if b.get('anchor_slug') == 'adct']

# Resumo
doc3.add_heading('Resumo', level=1)
data = [
    ('Blocos resolvidos', len(blocks)),
    ('Artigos cobertos', f'{len(arts_cobertos)} de 250'),
    ('Artigos sem decisao', len(faltando)),
    ('Art. 233', '0 (ADCT separado)'),
    ('Blocos ADCT', len(adct_blocks)),
    ('Blocos com processo', sum(1 for b in blocks if b.get('process_class'))),
    ('Blocos com relator', sum(1 for b in blocks if b.get('relator'))),
    ('Blocos com ano', sum(1 for b in blocks if b.get('ano'))),
    ('Relatores distintos', len(rels)),
    ('Classes distintas', len(cls_proc)),
    ('Blocos com inciso', sum(1 for b in blocks if b.get('anchor_inciso'))),
    ('Blocos com paragrafo', sum(1 for b in blocks if b.get('anchor_paragrafo'))),
    ('Blocos com alinea', sum(1 for b in blocks if b.get('anchor_alinea'))),
]
t = doc3.add_table(rows=len(data)+1, cols=2)
t.style = 'Light Grid Accent 1'
t.rows[0].cells[0].text = 'Metrica'
t.rows[0].cells[1].text = 'Valor'
for i, (k, v) in enumerate(data, 1):
    t.rows[i].cells[0].text = str(k)
    t.rows[i].cells[1].text = str(v)

doc3.add_heading('Artigos faltando', level=2)
doc3.add_paragraph(str(faltando))

doc3.add_heading('Top 30 artigos', level=1)
t2 = doc3.add_table(rows=31, cols=3)
t2.style = 'Light Grid Accent 1'
t2.rows[0].cells[0].text = 'Artigo'
t2.rows[0].cells[1].text = 'Blocos'
t2.rows[0].cells[2].text = 'Ramos'
for i, (art, n) in enumerate(por_artigo.most_common(30), 1):
    t2.rows[i].cells[0].text = f'Art. {art}'
    t2.rows[i].cells[1].text = str(n)
    t2.rows[i].cells[2].text = str(len(art_ramos.get(art, set())))

doc3.add_heading('Top 15 relatores', level=1)
t3 = doc3.add_table(rows=16, cols=2)
t3.style = 'Light Grid Accent 1'
t3.rows[0].cells[0].text = 'Relator'
t3.rows[0].cells[1].text = 'Blocos'
for i, (r, n) in enumerate(rels.most_common(15), 1):
    t3.rows[i].cells[0].text = r
    t3.rows[i].cells[1].text = str(n)

doc3.add_heading('Top 15 classes', level=1)
t4 = doc3.add_table(rows=16, cols=2)
t4.style = 'Light Grid Accent 1'
t4.rows[0].cells[0].text = 'Classe'
t4.rows[0].cells[1].text = 'Blocos'
for i, (c, n) in enumerate(cls_proc.most_common(15), 1):
    t4.rows[i].cells[0].text = c
    t4.rows[i].cells[1].text = str(n)

doc3.add_heading('Por editorial', level=1)
t5 = doc3.add_table(rows=len(por_editorial)+1, cols=2)
t5.style = 'Light Grid Accent 1'
t5.rows[0].cells[0].text = 'Editorial'
t5.rows[0].cells[1].text = 'Blocos'
for i, (ed, n) in enumerate(por_editorial.most_common(), 1):
    t5.rows[i].cells[0].text = ed
    t5.rows[i].cells[1].text = str(n)

doc3.save(os.path.join(BASE, "DIAGNOSTICO_PARSER_v8_final.docx"))
print("  Salvo.")

print("\n4 arquivos Word gerados:")
for f in ["ESPELHO_MISTO_v8_final.docx", "ESPELHO_CF_ancoragem_v8_final.docx",
          "ESPELHO_blocos_decisoes_v8_final.docx", "DIAGNOSTICO_PARSER_v8_final.docx"]:
    fpath = os.path.join(BASE, f)
    print(f"  {f}: {os.path.getsize(fpath)/1024:.0f} KB")
