from docx import Document
import json, re, csv
from collections import defaultdict, Counter

BASE = r'C:\projetos\icons\ingest'

# ══════════════════════════════════════════
# FASE A — Árvore canônica da CF pura
# ══════════════════════════════════════════
print('FASE A — Construindo árvore canônica...')
doc_cf = Document(f'{BASE}\\CF_1988_arvore_completa.docx')

norm_tree = []
ctx = {'titulo': None, 'capitulo': None, 'secao': None, 'subsecao': None}
node_id = 0

RE_ART  = re.compile(r'^Art\.?\s*(\d+)[º°]?\b')
RE_PAR  = re.compile(r'^§\s*(\d+|único)[º°]?\b', re.IGNORECASE)
RE_INC  = re.compile(r'^(I{1,3}V?|VI{0,3}|IX|X{1,3}I{0,3}|XL|L|XC|C)\s*[-–—]')
RE_ALI  = re.compile(r'^([a-z])\)\s')

current_art = current_par = current_inc = current_ali = None

for p in doc_cf.paragraphs:
    t = p.text.strip()
    if not t: continue
    s = p.style.name
    node_id += 1

    if s == 'Heading 1':
        ctx = {'titulo': t, 'capitulo': None, 'secao': None, 'subsecao': None}
        current_art = current_par = current_inc = current_ali = None
        norm_tree.append({'node_id': node_id, 'type': 'titulo', 'slug': f'cf-tit-{node_id}', 'label': t, 'parent': None, 'anchor_blocks': []})

    elif s == 'Heading 2':
        ctx['capitulo'] = t; ctx['secao'] = ctx['subsecao'] = None
        norm_tree.append({'node_id': node_id, 'type': 'capitulo', 'slug': f'cf-cap-{node_id}', 'label': t, 'parent': ctx['titulo'], 'anchor_blocks': []})

    elif s == 'Heading 3':
        ctx['secao'] = t; ctx['subsecao'] = None
        norm_tree.append({'node_id': node_id, 'type': 'secao', 'slug': f'cf-sec-{node_id}', 'label': t, 'parent': ctx['capitulo'], 'anchor_blocks': []})

    else:
        m = RE_ART.match(t)
        if m:
            current_art = int(m.group(1))
            current_par = current_inc = current_ali = None
            slug = f'cf-art-{current_art}'
            norm_tree.append({'node_id': node_id, 'type': 'artigo', 'slug': slug, 'number': current_art, 'label': t[:80], 'titulo': ctx['titulo'], 'capitulo': ctx['capitulo'], 'secao': ctx['secao'], 'parent': ctx['secao'] or ctx['capitulo'] or ctx['titulo'], 'anchor_blocks': []})
            continue

        m = RE_PAR.match(t)
        if m and current_art:
            current_par = m.group(1)
            current_inc = current_ali = None
            slug = f'cf-art-{current_art}-par-{current_par}'
            norm_tree.append({'node_id': node_id, 'type': 'paragrafo', 'slug': slug, 'number': current_par, 'label': t[:80], 'artigo': current_art, 'parent': f'cf-art-{current_art}', 'anchor_blocks': []})
            continue

        m = RE_INC.match(t)
        if m and current_art:
            current_inc = m.group(1)
            current_ali = None
            par_str = f'-par-{current_par}' if current_par else ''
            slug = f'cf-art-{current_art}{par_str}-inc-{current_inc}'
            norm_tree.append({'node_id': node_id, 'type': 'inciso', 'slug': slug, 'number': current_inc, 'label': t[:80], 'artigo': current_art, 'parent': f'cf-art-{current_art}{par_str}', 'anchor_blocks': []})
            continue

        m = RE_ALI.match(t)
        if m and current_art:
            current_ali = m.group(1)
            par_str = f'-par-{current_par}' if current_par else ''
            inc_str = f'-inc-{current_inc}' if current_inc else ''
            slug = f'cf-art-{current_art}{par_str}{inc_str}-ali-{current_ali}'
            norm_tree.append({'node_id': node_id, 'type': 'alinea', 'slug': slug, 'number': current_ali, 'label': t[:80], 'artigo': current_art, 'parent': f'cf-art-{current_art}{par_str}{inc_str}', 'anchor_blocks': []})

slug_index = {n['slug']: i for i, n in enumerate(norm_tree)}

# Build set of valid CF article numbers for cotejo
cf_valid_arts = set()
for n in norm_tree:
    if n['type'] == 'artigo':
        cf_valid_arts.add(n.get('number', 0))

print(f'  Nos na arvore: {len(norm_tree)}')
print(f'  Artigos validos CF: {len(cf_valid_arts)}')

# ══════════════════════════════════════════
# FASE B — Classificador A/B/C/D
# ══════════════════════════════════════════
print('\nFASE B — Classificando paragrafos do comentado...')
doc_main = Document(f'{BASE}\\cf_comentada.docx')

CLASSES_PROC = r'ADI|ADC|ADPF|RE|HC|MS|MI|RCL|ARE|RHC|ACO|PET|INQ|AP|RMS|STA|SS|Rcl|AgR|ED'
RE_META  = re.compile(rf'^\[.*({CLASSES_PROC}).*\]$', re.IGNORECASE)
RE_META2 = re.compile(rf'^\[?\s*({CLASSES_PROC})[\s\.\-]*\d', re.IGNORECASE)
RE_SUMV  = re.compile(r'^\[?\s*S[uú]mula\s+(Vinculante\s+)?n?[oº°]?\s*\d+', re.IGNORECASE)

RE_ART_ANCORA = re.compile(r'^Art\.?\s*(\d+)[º°]?\s+\S')

EDITORIAIS = [
    (re.compile(r'^Controle\s+concentrado', re.I),      'controle_concentrado'),
    (re.compile(r'^Repercuss[aã]o\s+geral', re.I),      'repercussao_geral'),
    (re.compile(r'^Julgados?\s+correlatos?', re.I),     'julgado_correlato'),
    (re.compile(r'^S[uú]mula\s+[Vv]inculante', re.I),  'sumula_vinculante'),
    (re.compile(r'^S[uú]mula\b', re.I),                 'sumula'),
]

def classifica_main(texto):
    t = texto.strip()
    # Classe A — âncora normativa
    # Art. só é âncora se o número pertence à CF (cotejo com árvore canônica)
    m_art = RE_ART_ANCORA.match(t)
    if m_art:
        art_num = int(m_art.group(1))
        if art_num in cf_valid_arts:
            return 'A'
        # Art. de outra lei — não é âncora, é comentário
    if RE_PAR.match(t) and len(t) > 5: return 'A'
    if RE_INC.match(t) and len(t) > 5: return 'A'
    if RE_ALI.match(t) and len(t) > 5: return 'A'
    for pat, _ in EDITORIAIS:
        if pat.match(t): return 'B'
    if RE_META.match(t) or RE_META2.match(t) or RE_SUMV.match(t): return 'D'
    return 'C'

def get_editorial(texto):
    for pat, cat in EDITORIAIS:
        if pat.match(texto.strip()):
            return cat
    return None

paragrafos_main = []
for i, p in enumerate(doc_main.paragraphs):
    t = p.text.strip()
    if not t: continue
    classe = classifica_main(t)
    paragrafos_main.append({'idx': i, 'classe': classe, 'texto': t})

dist = Counter(p['classe'] for p in paragrafos_main)
print(f'  Total paragrafos: {len(paragrafos_main)}')
print(f'  Distribuicao: {dict(dist)}')

# ══════════════════════════════════════════
# FASE C — Máquina de estados → blocos resolvidos
# ══════════════════════════════════════════
print('\nFASE C — Montando blocos resolvidos...')

state = {
    'artigo': None, 'paragrafo': None, 'inciso': None, 'alinea': None,
    'editorial': 'sem_categoria',
    'slug': None
}

def get_slug(state):
    art = state['artigo']
    if not art: return None
    par = state['paragrafo']
    inc = state['inciso']
    ali = state['alinea']
    slug = f'cf-art-{art}'
    if par: slug += f'-par-{par}'
    if inc: slug += f'-inc-{inc}'
    if ali: slug += f'-ali-{ali}'
    return slug

commentary_blocks = []
block_id = 0
current_block = None

def flush_block():
    global current_block
    if current_block and (current_block['block_text'].strip() or current_block['metadata_text']):
        commentary_blocks.append(current_block)
    current_block = None

def new_block():
    global block_id
    block_id += 1
    return {
        'block_id': block_id,
        'anchor_slug': state['slug'],
        'anchor_artigo': state['artigo'],
        'anchor_paragrafo': state['paragrafo'],
        'anchor_inciso': state['inciso'],
        'anchor_alinea': state['alinea'],
        'editorial_marker': state['editorial'],
        'block_text': '',
        'metadata_text': '',
        'paragraph_start': None,
        'paragraph_end': None
    }

for p in paragrafos_main:
    idx   = p['idx']
    cl    = p['classe']
    texto = p['texto']

    if cl == 'A':
        flush_block()
        m = RE_ART_ANCORA.match(texto)
        if m and int(m.group(1)) in cf_valid_arts:
            state = {'artigo': int(m.group(1)), 'paragrafo': None, 'inciso': None, 'alinea': None, 'editorial': 'sem_categoria', 'slug': None}
        elif RE_PAR.match(texto):
            mp = RE_PAR.match(texto)
            state['paragrafo'] = mp.group(1); state['inciso'] = None; state['alinea'] = None
        elif RE_INC.match(texto):
            mi = RE_INC.match(texto)
            state['inciso'] = mi.group(1); state['alinea'] = None
        elif RE_ALI.match(texto):
            ma = RE_ALI.match(texto)
            state['alinea'] = ma.group(1)
        state['slug'] = get_slug(state)
        if state['slug'] and state['slug'] in slug_index:
            norm_tree[slug_index[state['slug']]]['anchor_blocks'].append(block_id + 1)

    elif cl == 'B':
        flush_block()
        ed = get_editorial(texto)
        if ed: state['editorial'] = ed

    elif cl == 'C':
        if not current_block:
            current_block = new_block()
            current_block['paragraph_start'] = idx
        current_block['block_text'] += (' ' if current_block['block_text'] else '') + texto
        current_block['paragraph_end'] = idx

    elif cl == 'D':
        if not current_block:
            current_block = new_block()
            current_block['paragraph_start'] = idx
        current_block['metadata_text'] += (' ' if current_block['metadata_text'] else '') + texto
        current_block['paragraph_end'] = idx
        flush_block()

flush_block()
print(f'  Blocos resolvidos: {len(commentary_blocks)}')

# ══════════════════════════════════════════
# SAÍDAS
# ══════════════════════════════════════════

with open(f'{BASE}\\norm_tree.json', 'w', encoding='utf-8') as f:
    json.dump(norm_tree, f, ensure_ascii=False, indent=2)
print(f'\n> norm_tree.json — {len(norm_tree)} nos')

with open(f'{BASE}\\commentary_blocks.json', 'w', encoding='utf-8') as f:
    json.dump(commentary_blocks, f, ensure_ascii=False, indent=2)
print(f'> commentary_blocks.json — {len(commentary_blocks)} blocos')

with open(f'{BASE}\\commentary_blocks.tsv', 'w', encoding='utf-8-sig', newline='') as f:
    w = csv.DictWriter(f, fieldnames=['block_id','anchor_slug','anchor_artigo','anchor_paragrafo','anchor_inciso','anchor_alinea','editorial_marker','block_text','metadata_text','paragraph_start','paragraph_end'], delimiter='\t')
    w.writeheader()
    w.writerows(commentary_blocks)
print(f'> commentary_blocks.tsv')

# Diagnóstico final
print(f'\n=== DIAGNOSTICO FINAL ===')
print(f'Nos na arvore canonica: {len(norm_tree)}')
print(f'Blocos resolvidos: {len(commentary_blocks)}')

arts_cobertos = len(set(b['anchor_artigo'] for b in commentary_blocks if b['anchor_artigo']))
print(f'Artigos cobertos: {arts_cobertos}')

eds = Counter(b['editorial_marker'] for b in commentary_blocks)
print(f'Por marcador editorial:')
for ed, n in eds.most_common():
    print(f'  {ed}: {n}')

print(f'\nAmostra — Art. 5:')
art5 = [b for b in commentary_blocks if b['anchor_artigo'] == 5][:5]
for b in art5:
    print(f"  [{b['block_id']}] [{b['anchor_slug']}] [{b['editorial_marker']}]")
    print(f"    texto: {b['block_text'][:100]}")
    print(f"    meta:  {b['metadata_text'][:100]}")

print(f'\nAmostra — Art. 102:')
art102 = [b for b in commentary_blocks if b['anchor_artigo'] == 102][:5]
for b in art102:
    print(f"  [{b['block_id']}] [{b['anchor_slug']}] [{b['editorial_marker']}]")
    print(f"    texto: {b['block_text'][:100]}")
    print(f"    meta:  {b['metadata_text'][:100]}")
