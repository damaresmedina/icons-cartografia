from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json
import os

BASE = os.path.join("C:", os.sep, "projetos", "icons", "ingest")

# Load data
with open(os.path.join(BASE, "norm_tree.json"), "r", encoding="utf-8") as f:
    norm_tree = json.load(f)

with open(os.path.join(BASE, "commentary_blocks.json"), "r", encoding="utf-8") as f:
    blocks = json.load(f)

# Build index: artigo -> blocks
art_blocks = {}
for b in blocks:
    art = b.get("anchor_artigo")
    if art is not None:
        if art not in art_blocks:
            art_blocks[art] = []
        art_blocks[art].append(b)

# ══════════════════════════════════════════
# ARQUIVO 1: Constituição com ancoragem
# ══════════════════════════════════════════
print("Gerando arquivo 1: CF com ancoragem...")
doc1 = Document()

style = doc1.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(10)

doc1.add_heading('CONSTITUIÇÃO FEDERAL DE 1988', level=0)
doc1.add_heading('Árvore Normativa com Ancoragem de Decisões', level=1)
p = doc1.add_paragraph()
p.add_run(f'Nós normativos: {len(norm_tree)}').bold = True
p.add_run(f'  |  Blocos de decisão: {len(blocks)}')
doc1.add_paragraph('')

for node in norm_tree:
    nt = node['type']
    label = node.get('label', '')
    slug = node.get('slug', '')
    anchor_count = len(node.get('anchor_blocks', []))
    art_num = node.get('number', node.get('artigo', ''))

    if nt == 'titulo':
        doc1.add_heading(label, level=1)

    elif nt == 'capitulo':
        doc1.add_heading(label, level=2)

    elif nt == 'secao':
        doc1.add_heading(label, level=3)

    elif nt == 'artigo':
        p = doc1.add_paragraph()
        run = p.add_run(f'{label}')
        run.bold = True
        run.font.size = Pt(11)

        # Count decisions for this article
        dec_count = len(art_blocks.get(art_num, []))
        if dec_count > 0:
            run2 = p.add_run(f'  [{dec_count} decisões vinculadas]')
            run2.font.size = Pt(8)
            run2.font.color.rgb = RGBColor(0, 100, 180)

        p2 = doc1.add_paragraph()
        run3 = p2.add_run(f'slug: {slug}')
        run3.font.size = Pt(7)
        run3.font.color.rgb = RGBColor(128, 128, 128)

    elif nt == 'paragrafo':
        p = doc1.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        run = p.add_run(label)
        run.font.size = Pt(10)
        run4 = p.add_run(f'  [{slug}]')
        run4.font.size = Pt(7)
        run4.font.color.rgb = RGBColor(160, 160, 160)

    elif nt == 'inciso':
        p = doc1.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        run = p.add_run(label)
        run.font.size = Pt(10)

    elif nt == 'alinea':
        p = doc1.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.7)
        run = p.add_run(label)
        run.font.size = Pt(10)

out1 = os.path.join(BASE, "ESPELHO_1_CF_ancoragem.docx")
doc1.save(out1)
print(f"  Salvo: {out1}")

# ══════════════════════════════════════════
# ARQUIVO 2: Blocos de comentários/decisões
# ══════════════════════════════════════════
print("\nGerando arquivo 2: Blocos de comentários...")
doc2 = Document()

style2 = doc2.styles['Normal']
style2.font.name = 'Arial'
style2.font.size = Pt(10)

doc2.add_heading('BLOCOS DE DECISÕES E COMENTÁRIOS', level=0)
doc2.add_heading('Extraídos de "A Constituição e o Supremo"', level=1)
p = doc2.add_paragraph()
p.add_run(f'Total de blocos: {len(blocks)}').bold = True
doc2.add_paragraph('')

# Group by article
current_art = None
for b in blocks:
    art = b.get('anchor_artigo')
    slug = b.get('anchor_slug', 'sem-ancora')
    editorial = b.get('editorial_marker', 'sem_categoria')
    texto = b.get('block_text', '').strip()
    meta = b.get('metadata_text', '').strip()
    block_id = b.get('block_id', 0)

    # New article header
    if art != current_art:
        current_art = art
        if art:
            doc2.add_heading(f'Art. {art}', level=1)
        else:
            doc2.add_heading('Sem âncora normativa', level=1)

    # Editorial marker
    if editorial and editorial != 'sem_categoria':
        p_ed = doc2.add_paragraph()
        run_ed = p_ed.add_run(editorial.upper().replace('_', ' '))
        run_ed.bold = True
        run_ed.font.size = Pt(9)
        run_ed.font.color.rgb = RGBColor(180, 100, 0)

    # Block header
    p_header = doc2.add_paragraph()
    run_h = p_header.add_run(f'Bloco {block_id} | {slug} | {editorial}')
    run_h.font.size = Pt(7)
    run_h.font.color.rgb = RGBColor(128, 128, 128)

    # Block text
    if texto:
        p_texto = doc2.add_paragraph()
        p_texto.paragraph_format.left_indent = Inches(0.2)
        run_t = p_texto.add_run(texto[:1000])
        run_t.font.size = Pt(9)
        if len(texto) > 1000:
            run_more = p_texto.add_run(f' [...+{len(texto)-1000} chars]')
            run_more.font.size = Pt(7)
            run_more.font.color.rgb = RGBColor(160, 160, 160)

    # Metadata
    if meta:
        p_meta = doc2.add_paragraph()
        p_meta.paragraph_format.left_indent = Inches(0.2)
        run_m = p_meta.add_run(meta[:500])
        run_m.font.size = Pt(8)
        run_m.font.color.rgb = RGBColor(0, 80, 160)
        run_m.bold = True

    # Separator
    doc2.add_paragraph('─' * 60).runs[0].font.size = Pt(6)

out2 = os.path.join(BASE, "ESPELHO_2_blocos_decisoes.docx")
doc2.save(out2)
print(f"  Salvo: {out2}")

print(f"\nPronto.")
