import json, os, re
from collections import Counter, defaultdict
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_COLOR_INDEX

OUT = os.path.join("C:", os.sep, "projetos", "icons", "ingest", "saida_parser")
BASE = os.path.join("C:", os.sep, "projetos", "icons", "ingest")

with open(os.path.join(OUT, "norm_tree.json"), "r", encoding="utf-8") as f:
    norm_tree = json.load(f)
with open(os.path.join(OUT, "commentary_blocks.json"), "r", encoding="utf-8") as f:
    blocks = json.load(f)

# Hierarquia editorial
EDITORIAL_ORDER = {
    'sumula_vinculante': 1,
    'sumula': 2,
    'controle_concentrado': 3,
    'repercussao_geral': 4,
    'julgados_correlatos': 5,
    'julgado_correlato': 5,
    'sem_categoria': 6,
    '': 7,
}

def editorial_sort_key(b):
    ed = b.get('editorial_marker', '') or ''
    return EDITORIAL_ORDER.get(ed, 6)

# Index blocks by article, sorted by editorial hierarchy
blocks_by_art = defaultdict(list)
for b in blocks:
    art = b.get('anchor_artigo')
    if art:
        blocks_by_art[art].append(b)

for art in blocks_by_art:
    blocks_by_art[art].sort(key=editorial_sort_key)

# Colors:
# NORMA = azul escuro
# EDITORIAL = laranja
# DECISAO/COMENTARIO = preto
# METADADO = azul

AZUL_NORMA = RGBColor(0, 51, 102)
LARANJA_EDITORIAL = RGBColor(180, 90, 0)
PRETO_TEXTO = RGBColor(30, 30, 30)
AZUL_META = RGBColor(0, 60, 160)
CINZA_SLUG = RGBColor(140, 140, 140)

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(9)

doc.add_heading('CONSTITUICAO E O SUPREMO — ESPELHO MISTO', level=0)
p_sub = doc.add_paragraph()
p_sub.add_run('Norma em AZUL').font.color.rgb = AZUL_NORMA
p_sub.add_run(' | ')
p_sub.add_run('Editorial em LARANJA').font.color.rgb = LARANJA_EDITORIAL
p_sub.add_run(' | ')
p_sub.add_run('Decisao/comentario em PRETO').font.color.rgb = PRETO_TEXTO
p_sub.add_run(' | ')
r_m = p_sub.add_run('Metadado em AZUL NEGRITO')
r_m.font.color.rgb = AZUL_META
r_m.bold = True
total_blocos = sum(len(v) for v in blocks_by_art.values())
total_artigos = len(blocks_by_art)
doc.add_paragraph(f'{total_artigos} artigos cobertos | {total_blocos:,} blocos | Parser v9 (dedup + hierarquia editorial)')
doc.add_paragraph('')

# Walk the norm_tree and interleave blocks
current_art = None

for node in norm_tree:
    nt = node['type']
    label = node.get('label', '')
    art_num = node.get('number', node.get('artigo', ''))

    # NORMA — sempre em azul
    if nt == 'titulo':
        h = doc.add_heading(label, level=1)
        for run in h.runs:
            run.font.color.rgb = AZUL_NORMA

    elif nt == 'capitulo':
        h = doc.add_heading(label, level=2)
        for run in h.runs:
            run.font.color.rgb = AZUL_NORMA

    elif nt == 'secao':
        h = doc.add_heading(label, level=3)
        for run in h.runs:
            run.font.color.rgb = AZUL_NORMA

    elif nt == 'artigo':
        current_art = art_num
        dec_count = len(blocks_by_art.get(art_num, []))

        # Artigo em azul
        p = doc.add_paragraph()
        run = p.add_run(label)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = AZUL_NORMA

        if dec_count > 0:
            r2 = p.add_run(f'  [{dec_count} decisoes]')
            r2.font.size = Pt(7)
            r2.font.color.rgb = CINZA_SLUG

        # Insert all blocks for this article
        art_blocks = blocks_by_art.get(art_num, [])
        last_editorial = None

        for b in art_blocks:
            editorial = b.get('editorial_marker', '')
            slug = b.get('anchor_slug', '')
            texto = b.get('block_text', '').strip()
            meta = b.get('metadata_text', '').strip()
            proc = b.get('process_class', '')
            num = b.get('process_number', '')
            relator = b.get('relator', '')

            # Editorial marker in orange
            if editorial and editorial != 'sem_categoria' and editorial != last_editorial:
                last_editorial = editorial
                p_ed = doc.add_paragraph()
                p_ed.paragraph_format.left_indent = Inches(0.1)
                r_ed = p_ed.add_run(editorial.upper().replace('_', ' '))
                r_ed.bold = True
                r_ed.font.size = Pt(8)
                r_ed.font.color.rgb = LARANJA_EDITORIAL

            # Slug in gray
            p_slug = doc.add_paragraph()
            p_slug.paragraph_format.left_indent = Inches(0.2)
            r_slug = p_slug.add_run(f'{slug}')
            r_slug.font.size = Pt(6)
            r_slug.font.color.rgb = CINZA_SLUG
            if proc:
                r_proc = p_slug.add_run(f' | {proc} {num}')
                r_proc.font.size = Pt(6)
                r_proc.font.color.rgb = AZUL_META
            if relator:
                r_rel = p_slug.add_run(f' | rel. {relator}')
                r_rel.font.size = Pt(6)
                r_rel.font.color.rgb = AZUL_META

            # Text in black
            if texto:
                p_t = doc.add_paragraph()
                p_t.paragraph_format.left_indent = Inches(0.2)
                r_t = p_t.add_run(texto[:600])
                r_t.font.size = Pt(9)
                r_t.font.color.rgb = PRETO_TEXTO

            # Metadata in blue bold
            if meta:
                p_m = doc.add_paragraph()
                p_m.paragraph_format.left_indent = Inches(0.2)
                r_m = p_m.add_run(meta[:250])
                r_m.font.size = Pt(8)
                r_m.font.color.rgb = AZUL_META
                r_m.bold = True

    elif nt == 'paragrafo':
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        run = p.add_run(label)
        run.font.size = Pt(9)
        run.font.color.rgb = AZUL_NORMA

    elif nt == 'inciso':
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        run = p.add_run(label)
        run.font.size = Pt(9)
        run.font.color.rgb = AZUL_NORMA

    elif nt == 'alinea':
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.7)
        run = p.add_run(label)
        run.font.size = Pt(9)
        run.font.color.rgb = AZUL_NORMA

outpath = os.path.join(BASE, "ESPELHO_MISTO_v9.docx")
doc.save(outpath)
print(f"Salvo: {outpath}")
print(f"Nos normativos: {len(norm_tree)}")
print(f"Blocos inseridos: {sum(len(v) for v in blocks_by_art.values())}")
