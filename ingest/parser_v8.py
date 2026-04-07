import re
import json
import csv
import os
from collections import Counter, defaultdict
from docx import Document

BASE = os.path.join("C:", os.sep, "projetos", "icons", "ingest")
OUT = os.path.join(BASE, "saida_parser")
os.makedirs(OUT, exist_ok=True)

# ══════════════════════════════════════════
# EDITORIAIS + METADADOS
# ══════════════════════════════════════════

EDITORIAIS = [
    (re.compile(r'^Controle\s+concentrado', re.I),      'controle_concentrado'),
    (re.compile(r'^Controle\s+de\s+concentrado', re.I),  'controle_concentrado'),
    (re.compile(r'^Repercuss[aã]o\s+geral', re.I),      'repercussao_geral'),
    (re.compile(r'^Julgados?\s+[Cc]orrelatos?', re.I),  'julgados_correlatos'),
    (re.compile(r'^S[uú]mulas?\s+[Vv]inculantes?', re.I), 'sumula_vinculante'),
    (re.compile(r'^S[uú]mulas?\s*$', re.I),             'sumula'),
]

CLASSES_PROC = r'ADI|ADC|ADPF|ADO|RE|ARE|HC|RHC|MS|MI|AP|ACO|Rcl|AI|Pet|SS|SL|IF|Inq|Ext|AC|AgR|ED|MC|QO|SV|PSV|AO|AR|RMS|CR|Rp|STA'
RE_META  = re.compile(rf'^\[.*({CLASSES_PROC}).*\]$', re.IGNORECASE)
RE_META2 = re.compile(rf'^\[?\s*({CLASSES_PROC})[\s\.\-]*\d', re.IGNORECASE)
RE_SUMV  = re.compile(r'^\[?\s*S[uú]mula\s+(Vinculante\s+)?n?[oº°]?\s*\d+', re.IGNORECASE)

RE_PROC   = re.compile(rf'\b({CLASSES_PROC})\s+[\-n.]*\s*([\d.]+(?:[-/]\d+)?)', re.IGNORECASE)
RE_REL    = re.compile(r'rel\.?\s*(?:p/\s*o?\s*ac\.?\s*)?min\.?\s*([^,\]]+)', re.IGNORECASE)
RE_JULG   = re.compile(r'j\.\s*(\d[\d\-]+(?:\s*de\s*\w+\s*de\s*\d+)?)', re.IGNORECASE)
RE_DJE    = re.compile(r'D[Jj][Ee]?\s*de\s*(\d[\d\-/]+)', re.IGNORECASE)
RE_TEMA   = re.compile(r'[Tt]ema\s*(\d+)', re.IGNORECASE)

RE_INC = re.compile(r'^(I{1,3}V?|VI{0,3}|IX|X{1,3}I{0,3}|XL|L|XC|C{1,3}|D)\s*[-\u2013\u2014]\s*\S')
RE_PAR = re.compile(r'^§\s*(\d+|[uú]nico)[º°]?\s', re.IGNORECASE)
RE_ALI = re.compile(r'^([a-z])\)\s')

# ══════════════════════════════════════════
# FASE A — Árvore canônica
# ══════════════════════════════════════════

print("FASE A — Arvore canonica da CF...")
doc_cf = Document(os.path.join(BASE, "CF_1988_arvore_completa.docx"))

RE_ART_CF = re.compile(r'^Art\.?\s*(\d+)[º°]?\b')
RE_PAR_CF = re.compile(r'^§\s*(\d+|único)[º°]?\b', re.IGNORECASE)
RE_INC_CF = re.compile(r'^(I{1,3}V?|VI{0,3}|IX|X{1,3}I{0,3}|XL|L|XC|C{1,3}|D)\s*[-\u2013\u2014]')
RE_ALI_CF = re.compile(r'^([a-z])\)\s')

norm_tree = []
ctx = {'titulo': None, 'capitulo': None, 'secao': None}
cur_art = cur_par = cur_inc = None
nid = 0

for p in doc_cf.paragraphs:
    t = p.text.strip()
    if not t: continue
    s = p.style.name
    nid += 1

    if s == 'Heading 1':
        ctx = {'titulo': t, 'capitulo': None, 'secao': None}
        cur_art = cur_par = cur_inc = None
        norm_tree.append({'nid': nid, 'type': 'titulo', 'label': t, 'parent': None})
    elif s == 'Heading 2':
        ctx['capitulo'] = t; ctx['secao'] = None
        norm_tree.append({'nid': nid, 'type': 'capitulo', 'label': t, 'parent': ctx['titulo']})
    elif s == 'Heading 3':
        ctx['secao'] = t
        norm_tree.append({'nid': nid, 'type': 'secao', 'label': t, 'parent': ctx['capitulo']})
    else:
        m = RE_ART_CF.match(t)
        if m:
            cur_art = int(m.group(1)); cur_par = cur_inc = None
            norm_tree.append({'nid': nid, 'type': 'artigo', 'number': cur_art, 'label': t[:80],
                              'titulo': ctx['titulo'], 'capitulo': ctx['capitulo'], 'secao': ctx['secao']})
            continue
        m = RE_PAR_CF.match(t)
        if m and cur_art:
            cur_par = m.group(1); cur_inc = None
            norm_tree.append({'nid': nid, 'type': 'paragrafo', 'number': cur_par, 'artigo': cur_art, 'label': t[:80]})
            continue
        m = RE_INC_CF.match(t)
        if m and cur_art:
            cur_inc = m.group(1)
            norm_tree.append({'nid': nid, 'type': 'inciso', 'number': cur_inc, 'artigo': cur_art, 'label': t[:80]})
            continue
        m = RE_ALI_CF.match(t)
        if m and cur_art:
            norm_tree.append({'nid': nid, 'type': 'alinea', 'number': m.group(1), 'artigo': cur_art, 'label': t[:80]})

cf_valid_arts = set(n['number'] for n in norm_tree if n['type'] == 'artigo')

incisos_validos   = defaultdict(set)
paragrafos_validos = defaultdict(set)
alineas_validas   = defaultdict(set)

for no in norm_tree:
    art = no.get('artigo')
    if not art: continue
    if no.get('type') == 'inciso':
        incisos_validos[art].add(no['number'])
    elif no.get('type') == 'paragrafo':
        paragrafos_validos[art].add(str(no['number']))
    elif no.get('type') == 'alinea':
        alineas_validas[art].add(no['number'])

print(f"  Nos: {len(norm_tree)}")
print(f"  Artigos validos: {len(cf_valid_arts)}")

# ══════════════════════════════════════════
# PARTE 1 — Pré-computa mapa de âncoras
# ══════════════════════════════════════════

print("\nFASE B — Pre-computando mapa de ancoras...")

# Extract first 40 chars of article body from CF tree
art_corpo = {}
for p in doc_cf.paragraphs:
    t = p.text.strip()
    m = re.match(r'^Art\.?\s*(\d+)[º°]?\s+(.*)', t)
    if m:
        num = int(m.group(1))
        corpo = m.group(2).strip()[:40]
        if 1 <= num <= 250 and num not in art_corpo:
            art_corpo[num] = corpo

# Read main document paragraphs
doc_main = Document(os.path.join(BASE, "cf_comentada.docx"))
pars_main = []
for i, p in enumerate(doc_main.paragraphs):
    t = p.text.strip()
    if t:
        pars_main.append((i, t))

# Detect ADCT
ADCT_INICIO = None
for idx, t in pars_main:
    if re.match(r'^ATO DAS DISPOSI', t, re.IGNORECASE) and len(t) < 100 and idx > 20000:
        ADCT_INICIO = idx
        break
print(f"  ADCT detectado na posicao: {ADCT_INICIO}")

# Build ancora_map: art_num -> paragraph_index in main doc
ancora_map = {}
for num in range(1, 251):
    corpo = art_corpo.get(num, '')
    if corpo:
        fp = re.sub(r'\s+', ' ', corpo[:20]).strip().lower()
        for idx, t in pars_main:
            if ADCT_INICIO and idx > ADCT_INICIO:
                break
            m = re.match(r'^Art\.?\s*(\d+)[º°]?', t)
            if m and int(m.group(1)) == num:
                resto = re.sub(r'\s+', ' ', t[m.end():].strip()[:20]).lower()
                if fp[:10] in resto or resto[:10] in fp:
                    ancora_map[num] = idx
                    break

    # Fallback: first occurrence matching Art. N before ADCT
    if num not in ancora_map:
        for idx, t in pars_main:
            if ADCT_INICIO and idx > ADCT_INICIO:
                break
            m = re.match(r'^Art\.?\s*(\d+)[º°]?[\s\.]', t)
            if m and int(m.group(1)) == num:
                ancora_map[num] = idx
                break

pos_to_art = {v: k for k, v in ancora_map.items()}
print(f"  Ancoras mapeadas: {len(ancora_map)}")
print(f"  Artigos sem ancora: {sorted(set(range(1,251)) - set(ancora_map.keys()))}")

# ══════════════════════════════════════════
# FASE C — Classificador + Máquina de estados
# ══════════════════════════════════════════

print("\nFASE C — Maquina de estados...")

def build_slug(state):
    art = state['artigo']
    if not art: return None
    slug = f'cf-art-{art}'
    if state['paragrafo']: slug += f'-par-{state["paragrafo"]}'
    if state['inciso']:    slug += f'-inc-{state["inciso"]}'
    if state['alinea']:    slug += f'-ali-{state["alinea"]}'
    return slug

def extract_meta(text):
    meta = {'process_class': '', 'process_number': '', 'relator': '', 'data_julgamento': '', 'dje': '', 'tema': ''}
    m = RE_PROC.search(text)
    if m: meta['process_class'] = m.group(1).upper(); meta['process_number'] = m.group(2)
    m = RE_REL.search(text)
    if m: meta['relator'] = m.group(1).strip()
    m = RE_JULG.search(text)
    if m: meta['data_julgamento'] = m.group(1).strip()
    m = RE_DJE.search(text)
    if m: meta['dje'] = m.group(1).strip()
    m = RE_TEMA.search(text)
    if m: meta['tema'] = m.group(1)
    return meta

state = {'artigo': None, 'paragrafo': None, 'inciso': None, 'alinea': None, 'editorial': 'sem_categoria', 'slug': None}
blocks = []
block_id = 0
current_block = None
rejected_subnodes = 0

def flush():
    global current_block
    if current_block and (current_block['block_text'].strip() or current_block['metadata_text'].strip()):
        merged = (current_block['block_text'] + ' ' + current_block['metadata_text']).strip()
        meta = extract_meta(merged)
        current_block.update(meta)
        blocks.append(current_block)
    current_block = None

def new_block(idx):
    global block_id
    block_id += 1
    return {
        'block_id': block_id,
        'anchor_slug': state['slug'],
        'anchor_artigo': state['artigo'],
        'anchor_paragrafo': state['paragrafo'],
        'anchor_inciso': state['inciso'],
        'anchor_alinea': state['alinea'],
        'editorial_marker': state['editorial'],
        'block_text': '',
        'metadata_text': '',
        'paragraph_start': idx,
        'paragraph_end': idx
    }

def add_text(idx, texto):
    global current_block
    if not current_block:
        current_block = new_block(idx)
    current_block['block_text'] += (' ' if current_block['block_text'] else '') + texto
    current_block['paragraph_end'] = idx

# Process each paragraph
for par_idx, texto in pars_main:
    t = texto.strip()

    # 1. Check if this is a pre-computed article anchor
    if par_idx in pos_to_art:
        flush()
        num = pos_to_art[par_idx]
        state = {'artigo': num, 'paragrafo': None, 'inciso': None, 'alinea': None,
                 'editorial': 'sem_categoria', 'slug': f'cf-art-{num}'}
        continue

    # 2. Parágrafo normativo
    m = RE_PAR.match(t)
    if m and len(t) < 300 and state['artigo']:
        par_num = m.group(1).lower().replace('\u00fa', 'u')
        if par_num in paragrafos_validos.get(state['artigo'], set()) or not paragrafos_validos.get(state['artigo']):
            flush()
            state['paragrafo'] = par_num
            state['inciso'] = None
            state['alinea'] = None
            state['slug'] = build_slug(state)
        else:
            rejected_subnodes += 1
            add_text(par_idx, texto)
        continue

    # Parágrafo único alternativo
    m = re.match(r'^Par[a\u00e1]grafo\s+[u\u00fa]nico', t, re.IGNORECASE)
    if m and len(t) < 300 and state['artigo']:
        if 'unico' in paragrafos_validos.get(state['artigo'], set()) or not paragrafos_validos.get(state['artigo']):
            flush()
            state['paragrafo'] = 'unico'
            state['inciso'] = None
            state['alinea'] = None
            state['slug'] = build_slug(state)
        else:
            rejected_subnodes += 1
            add_text(par_idx, texto)
        continue

    # 3. Inciso normativo
    m = RE_INC.match(t)
    if m and len(t) < 200 and state['artigo']:
        inc_num = m.group(1)
        if inc_num in incisos_validos.get(state['artigo'], set()) or not incisos_validos.get(state['artigo']):
            flush()
            state['inciso'] = inc_num
            state['alinea'] = None
            state['slug'] = build_slug(state)
        else:
            rejected_subnodes += 1
            add_text(par_idx, texto)
        continue

    # 4. Alínea normativa
    m = RE_ALI.match(t)
    if m and len(t) < 200 and state['artigo']:
        ali_num = m.group(1)
        if ali_num in alineas_validas.get(state['artigo'], set()) or not alineas_validas.get(state['artigo']):
            flush()
            state['alinea'] = ali_num
            state['slug'] = build_slug(state)
        else:
            rejected_subnodes += 1
            add_text(par_idx, texto)
        continue

    # 5. Editorial
    for pat, cat in EDITORIAIS:
        if pat.match(t):
            flush()
            state['editorial'] = cat
            break
    else:
        # 6. Metadado
        if RE_META.match(t) or RE_META2.match(t) or RE_SUMV.match(t):
            if not current_block:
                current_block = new_block(par_idx)
            current_block['metadata_text'] += (' ' if current_block['metadata_text'] else '') + texto
            current_block['paragraph_end'] = par_idx
            flush()
        else:
            # 7. Texto
            add_text(par_idx, texto)

flush()
print(f"  Blocos: {len(blocks)}")
print(f"  Sub-nos rejeitados: {rejected_subnodes}")

# ══════════════════════════════════════════
# DIAGNÓSTICO
# ══════════════════════════════════════════

arts_cobertos = set(b['anchor_artigo'] for b in blocks if b['anchor_artigo'])
faltando = sorted(set(range(1, 251)) - arts_cobertos)

por_editorial = Counter(b['editorial_marker'] for b in blocks)
por_artigo = Counter(b['anchor_artigo'] for b in blocks if b['anchor_artigo'])

art_ramos = defaultdict(set)
for b in blocks:
    if b.get('anchor_artigo') and b.get('anchor_slug'):
        art_ramos[b['anchor_artigo']].add(b['anchor_slug'])

rels = Counter(b['relator'] for b in blocks if b.get('relator'))

print(f"\n=== DIAGNOSTICO FINAL ===")
print(f"  Nos arvore: {len(norm_tree)}")
print(f"  Ancoras mapeadas: {len(ancora_map)}")
print(f"  Blocos: {len(blocks)}")
print(f"  Artigos cobertos: {len(arts_cobertos)}")
print(f"  Faltando ({len(faltando)}): {faltando}")
print(f"  Sub-nos rejeitados: {rejected_subnodes}")

print(f"\n  Blocos com inciso: {sum(1 for b in blocks if b.get('anchor_inciso'))}")
print(f"  Blocos com paragrafo: {sum(1 for b in blocks if b.get('anchor_paragrafo'))}")
print(f"  Blocos com alinea: {sum(1 for b in blocks if b.get('anchor_alinea'))}")

slug_counts = Counter(b['anchor_slug'] for b in blocks if b['anchor_slug'])
print(f"\n  TOP 20 slugs:")
for slug, cnt in slug_counts.most_common(20):
    print(f"    {slug}: {cnt}")

print(f"\n  TOP 20 artigos:")
for art, n in por_artigo.most_common(20):
    print(f"    Art.{art}: {n}")

print(f"\n  Top 10 artigos com mais ramos:")
for art, ramos in sorted(art_ramos.items(), key=lambda x: len(x[1]), reverse=True)[:10]:
    print(f"    Art.{art}: {len(ramos)} ramos")

art5 = [b for b in blocks if b['anchor_artigo'] == 5]
inc5 = Counter(b.get('anchor_inciso') or 'raiz' for b in art5)
print(f"\n  Art. 5: {len(art5)} blocos")
for inc, cnt in inc5.most_common(5):
    print(f"    inc-{inc}: {cnt}")

art37 = [b for b in blocks if b['anchor_artigo'] == 37]
print(f"  Art. 37: {len(art37)} blocos")

art102 = [b for b in blocks if b['anchor_artigo'] == 102]
print(f"  Art. 102: {len(art102)} blocos")

art233 = [b for b in blocks if b['anchor_artigo'] == 233]
print(f"  Art. 233: {len(art233)} blocos")

print(f"\n  Por editorial:")
for ed, n in por_editorial.most_common():
    print(f"    {ed}: {n}")

procs = sum(1 for b in blocks if b['process_class'])
print(f"\n  Blocos com processo: {procs}")
print(f"  Top 5 relatores:")
for r, n in rels.most_common(5):
    print(f"    {r}: {n}")

# ══════════════════════════════════════════
# SALVAR
# ══════════════════════════════════════════

with open(os.path.join(OUT, 'norm_tree.json'), 'w', encoding='utf-8') as f:
    json.dump(norm_tree, f, ensure_ascii=False, indent=2)
with open(os.path.join(OUT, 'commentary_blocks.json'), 'w', encoding='utf-8') as f:
    json.dump(blocks, f, ensure_ascii=False, indent=2)
with open(os.path.join(OUT, 'commentary_blocks.tsv'), 'w', encoding='utf-8-sig', newline='') as f:
    if blocks:
        w = csv.DictWriter(f, fieldnames=list(blocks[0].keys()), delimiter='\t')
        w.writeheader()
        w.writerows(blocks)

for fname in ["norm_tree.json", "commentary_blocks.json", "commentary_blocks.tsv"]:
    fpath = os.path.join(OUT, fname)
    print(f"\n  {fname}: {os.path.getsize(fpath)/1024:.0f} KB")
